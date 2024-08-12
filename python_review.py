# 1. READING & WRITING TO A TEXT FILE
# 1.1: Read & print names
def read_txt_file(file):
    with open(file, 'r') as names_file:
        print(names_file.read())
# read_txt_file('names.txt')

# 1.2: Append new lines
def append_txt_file(file):
    with open(file, 'a') as names_file:
        names_file.write('\nPathik Bhattarai')
# append_txt_file('names.txt')


# 2. READING & WRITING TO A DOC FILE
from docx import Document
# 2.1: Extract all headings
def extract_docx_headings(file):
    doc_file = Document(file)
    for doc_para in doc_file.paragraphs:
        if doc_para.style.name.startswith('Heading'):
            print(doc_para.text)
# extract_docx_headings('headings.docx')

# 2.2: Create table containing student names and grades
student_details = [
    ('Name', 'Grade'),
    ('Pathik Bhattarai', 'A+'),
    ('Lionel Messi', 'A'),
    ('Lamine Yamal', 'B'),
    ]
def create_docx_table(file):
    doc_file = Document()
    doc_file.add_heading('Final Grades', 0)
    doc_table = doc_file.add_table(rows=len(student_details), cols=2)
    
    for i, student in enumerate(student_details):
        doc_table.cell(i, 0).text = student[0]
        doc_table.cell(i, 1).text = student[1]

    doc_file.save(file)
# create_docx_table('grades.docx')


# 3. READING & WRITING TO AN EXCEL FILE
from openpyxl import load_workbook, Workbook
# 3.1: Read & print data
def read_xlsx_file(file):
    wb = load_workbook(file)
    ws = wb.active
    print(list(row for row in ws.iter_rows(values_only=True)))
# read_xlsx_file('data.xlsx')

# 3.2: Create & write data
product_details = [
    ('Product', 'Quantity', 'Price'),
    ('Smartphone', 15, 800.00),
    ('Television', 25, 599.00),
    ('Laptop', 15, 999.99),
]
def create_xlsx_file(file):
    wb = Workbook()
    ws = wb.active
    for product in product_details:
        ws.append(product)
    wb.save(file)
# create_xlsx_file('products.xlsx')


# 4. MAKING HTTP CALLS
import requests
# 4.1: Fetch & display data from API
def fetch_api_data(url):
    response = requests.get(url)
    print(response.json())
# fetch_api_data('https://jsonplaceholder.typicode.com/comments')

# 4.2: program to submit a form using post request
form_details = {
    'name': 'Pathik Bhattarai',
    'age': 20,
    'email': 'pathik.b45@gmail.com',
}
def submit_form(url, data):
    r=requests.post(url, data)
# submit_form('https://httpbin.org/post', form_details)


# 5. ASYNCHRONOUS PROGRAMMING
import asyncio
import time
# 5.1: Print after waiting 3 seconds
async def print_after_wait():
    start_time = time.time()
    print(f'Starting program...')
    await asyncio.sleep(3)
    end_time = time.time()
    print(f'Program finished after {end_time - start_time} seconds')
# asyncio.run(print_after_wait())

# 5.2: Run multiple concurrent tasks and print completion time
async def run_concurrent_tasks():
    all_tasks = [print_after_wait() for _ in range(3)]
    await asyncio.gather(*all_tasks)
# asyncio.run(run_concurrent_tasks())


# 6. WEB SCRAPING
from bs4 import BeautifulSoup
# 6.1: Extract titles from blog posts
def extract_blog_titles(url):
    response = requests.get(url)
    soup = BeautifulSoup(response.content, 'html.parser')
    titles = soup.find_all('h2', class_='title is-5')
    print([title.text for title in titles])
# extract_blog_titles('https://realpython.github.io/fake-jobs/')

# 6.2: Extract all links from a webpage
def extract_links(url):
    response = requests.get(url)
    soup = BeautifulSoup(response.content, 'html.parser')
    links = soup.find_all('a')
    print([link.get('href') for link in links])
# extract_links('https://realpython.github.io/fake-jobs/')


# 7. DISCORD BOT
import os
from dotenv import load_dotenv
import discord
import requests

load_dotenv()
TOKEN = os.getenv('TOKEN')
GUILD = os.getenv('GUILD')

intents = discord.Intents.default()
intents.message_content = True
client = discord.Client(intents=intents)

# 7.1: Bot that echoes user messages
@client.event
async def on_ready():
    print(f'Bot is connected.')

@client.event
async def on_message(message):
    if message.author == client.user:
        return
    await message.channel.send(message.content)
# client.run(TOKEN)

# 7.2: Bot that responds with a joke
@client.event
async def on_ready():
    print('Bot is connected.')
    
@client.event
async def on_message(message):
    if message.content.startswith('!joke'):
        joke = get_jokes('https://official-joke-api.appspot.com/random_joke')
        if joke:
            await message.channel.send(f'{joke.get("setup")}\n{joke.get("punchline")}')
        else:
            await message.channel.send('No jokes available!')

def get_jokes(url):
    response = requests.get(url)
    return response.json()  
# client.run(TOKEN)


# 8. MULTITHREADING & MULTIPROCESSING
import threading
import multiprocessing

# 8.1: Download multiple images using multithreading
def multithreaded_download():
    def download_image(url, filename):
        response = requests.get(url)
        with open(filename, 'wb') as file:
            file.write(response.content)

    image_links = [
        'https://picsum.photos/200',
        'https://picsum.photos/200',
        'https://picsum.photos/200',
        'https://picsum.photos/200',
        'https://picsum.photos/200',
    ]

    def download_images():
        for i, link in enumerate(image_links):
            thread = threading.Thread(target=download_image, args=(link, f'image_{i}.jpg'))
            thread.start()
    download_images()
# multithreaded_download()

# 8.2: Create a multiprocessing program to calculate the sum of numbers in a large list by dividing the task among multiple processes.
def add_numbers(numbers, result, index):
    result[index] = sum(numbers)

def multiprocessed_addition():
    numbers = list(range(1000000))
    processes = []
    chunk_size = len(numbers) // 4
    manager = multiprocessing.Manager()
    result = manager.dict()

    for i in range(4):
        start = i * chunk_size
        end = (i + 1) * chunk_size
        process = multiprocessing.Process(target=add_numbers, args=(numbers[start:end], result, i))
        processes.append(process)
        process.start()

    for process in processes:
        process.join()

    total_sum = sum(result.values())
    print(f"Total sum: {total_sum}")

# if __name__ == '__main__':
#     multiprocessed_addition()
